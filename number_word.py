
#!/usr/bin/python3
# -*- coding: utf-8 -*-
import docx
import random
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import *
from docx.oxml.ns import qn

'''
#docx模块官方网站，可查询模块和函数用法
http://python-docx.readthedocs.io/en/latest/
pip install python-docx  
'''
#定义多少之内的加减法
zhilei=100

#定义生成的页数
NUMOFPAGE=20

#生成一个word对象file
file=docx.Document()

#设置页面的左右边界和上下边界
for section in file.sections:
    section.left_margin=Inches(0.8)
    section.right_margin = Inches(0.8)
    section.up_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)


#k循环用于产生对应页数
for k in range(NUMOFPAGE):
    #增加每一页的标题
    para = file.add_paragraph()
    run = para.add_run("天天算一算,练成大本领!("+str(zhilei)+"以内的加减法) \n    姓名：                      得分:                     日期：           ")

    #下面两行用于设置字体和字号
    run.font.name = u"微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u"微软雅黑")
    run.font.size = Pt(15)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #生成一个20行*5列的表格，每一个单元格中放一个问题等式。表格的作用主要用于排版
    table = file.add_table(rows=20, cols=5)
    table.style.font.name = "Arial"
    table.style.font.size = Pt(12)

    for i in range(20):
        print('.',end=' ')#这一行为测试输出的代码
        for j in range(5):
            #随机生成一个加减符号
            op=random.choice(["+","-"])
            #随机生成两个操作数
            op1=int(random.randint(1,zhilei-1))
            op2=int(random.randint(1,zhilei-1))

            if op=="+":
                op2=int(random.randint(1,(zhilei-op1)))#防止两个数的和超过zhilei
                run=table.cell(i,j).paragraphs[0].add_run(str(op1)+op+str(op2)+"=")
                run.font.size=Pt(14)
                table.cell(i,j).paragraphs[0].alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                if op1>op2:
                    run = table.cell(i, j).paragraphs[0].add_run(str(op1) + op + str(op2) + "=")
                    run.font.size = Pt(14)
                else:
                    run = table.cell(i, j).paragraphs[0].add_run(str(op2) + op + str(op1) + "=")
                    run.font.size = Pt(14)
    file.add_page_break()

    print('OK,第{0}页已经生成'.format(k+1))
    
#保存文件
file.save("mysonmath.docx")